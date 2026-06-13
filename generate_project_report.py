from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from pathlib import Path

ROOT = Path(__file__).resolve().parent
OUTPUT = ROOT / "Care-Grid_Project_Report_Parth_Kadiya.docx"
LOGO = ROOT / "report_assets" / "image2.png"

NAVY = "173B63"
TEAL = "0F8C77"
LIGHT_TEAL = "E8F5F2"
LIGHT_BLUE = "EAF1F8"
LIGHT_GRAY = "F3F5F7"
DARK = RGBColor(31, 45, 61)
MUTED = RGBColor(90, 105, 120)
WHITE = RGBColor(255, 255, 255)

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.27)
section.page_height = Inches(11.69)
section.top_margin = Inches(0.85)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(10.5)
normal.font.color.rgb = DARK
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.15
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for name, size, color, before, after in [
    ("Heading 1", 17, NAVY, 16, 8),
    ("Heading 2", 14, TEAL, 12, 6),
    ("Heading 3", 12, NAVY, 9, 4),
]:
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

caption = styles["Caption"]
caption.font.name = "Calibri"
caption.font.size = Pt(9)
caption.font.italic = True
caption.font.color.rgb = MUTED
caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.paragraph_format.space_after = Pt(8)

if "Code Block" not in [s.name for s in styles]:
    code_style = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
else:
    code_style = styles["Code Block"]
code_style.font.name = "Consolas"
code_style.font.size = Pt(8.5)
code_style.font.color.rgb = RGBColor(30, 41, 59)
code_style.paragraph_format.left_indent = Inches(0.2)
code_style.paragraph_format.right_indent = Inches(0.2)
code_style.paragraph_format.space_before = Pt(4)
code_style.paragraph_format.space_after = Pt(7)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="CBD5E1", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        elem = borders.find(qn(tag))
        if elem is None:
            elem = OxmlElement(tag)
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:color"), color)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(9)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_header_footer(sec, label="CARE-GRID | PROJECT REPORT"):
    header = sec.header
    p = header.paragraphs[0]
    p.text = label
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].font.size = Pt(8.5)
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = MUTED
    footer_p = sec.footer.paragraphs[0]
    footer_p.text = ""
    add_page_number(footer_p)


configure_header_footer(section)


def add_center(text, size=11, bold=False, color=DARK, after=6, italic=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.name = "Calibri"
    r.font.size = Pt(size)
    r.font.color.rgb = color
    return p


def add_para(text, bold_start=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    if bold_start and text.startswith(bold_start):
        r = p.add_run(bold_start)
        r.bold = True
        p.add_run(text[len(bold_start):])
    else:
        p.add_run(text)
    return p


def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p.add_run(item)


def add_numbers(items):
    numbering = doc.part.numbering_part.element
    abstract_id = None
    for abstract in numbering.findall(qn("w:abstractNum")):
        for p_style in abstract.findall(".//" + qn("w:pStyle")):
            if p_style.get(qn("w:val")) == "ListNumber":
                abstract_id = abstract.get(qn("w:abstractNumId"))
                break
        if abstract_id is not None:
            break
    if abstract_id is None:
        abstract_id = "0"
    existing_ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    num_id = str(max(existing_ids, default=0) + 1)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), abstract_id)
    num.append(abstract_ref)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.18)
        p_pr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), num_id)
        num_pr.extend([ilvl, num_id_node])
        p_pr.append(num_pr)
        p.add_run(item)


def add_table(headers, rows, widths=None, header_fill=NAVY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0]
    repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(9)
        if widths:
            cell.width = Inches(widths[i])
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cell = cells[i]
            if ri % 2 == 1:
                set_cell_shading(cell, "F8FAFC")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if len(str(text)) > 25 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(text))
            r.font.size = Pt(8.7)
            if widths:
                cell.width = Inches(widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_caption(text):
    doc.add_paragraph(text, style="Caption")


def add_callout(title, text, fill=LIGHT_TEAL):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=TEAL, size="8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 140, 180, 140, 180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(TEAL)
    p.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_code(lines):
    p = doc.add_paragraph(style="Code Block")
    p.add_run(lines)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_GRAY)
    p_pr.append(shd)


def add_chapter(title, subtitle):
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(110)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string(NAVY)
    add_center(subtitle, 12, False, RGBColor.from_string(TEAL), after=20)
    add_center("CARE-GRID: DOCTOR WEBSITE MAKER", 9, True, MUTED, after=6)
    doc.add_page_break()


def add_architecture_figure():
    table = doc.add_table(rows=7, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    labels = [
        ("CLIENT LAYER", "React 19 + Vite + Tailwind CSS | Doctor website, doctor admin, super-admin"),
        ("↓ HTTPS / REST API", ""),
        ("APPLICATION LAYER", "Node.js + Express | Validation, authentication, appointments, chatbot"),
        ("↓ ODM / Integrations", ""),
        ("DATA & SERVICE LAYER", "MongoDB Atlas | Cloudinary | Gmail SMTP | OpenAI API"),
        ("↓ Deployment", ""),
        ("PRODUCTION PLATFORM", "Vercel frontend + Render backend + MongoDB Atlas database"),
    ]
    for i, (head, detail) in enumerate(labels):
        cell = table.cell(i, 0)
        set_cell_margins(cell, 130, 160, 130, 160)
        cell.width = Inches(5.8)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        if "↓" in head:
            r = p.add_run(head)
            r.bold = True
            r.font.color.rgb = RGBColor.from_string(TEAL)
        else:
            set_cell_shading(cell, NAVY if i in (0, 2, 4, 6) else LIGHT_BLUE)
            r = p.add_run(head + "\n")
            r.bold = True
            r.font.color.rgb = WHITE
            r.font.size = Pt(10)
            r2 = p.add_run(detail)
            r2.font.color.rgb = WHITE
            r2.font.size = Pt(8.5)
    set_table_borders(table, color="FFFFFF", size="10")
    add_caption("Figure 3.1: High-level Care-Grid system architecture")


def add_workflow_figure():
    add_table(
        ["Step", "Doctor onboarding workflow", "System action"],
        [
            ("1", "Fill website creation form", "Validate structured doctor, clinic and media data"),
            ("2", "Verify registered email using OTP", "Issue purpose-bound, hashed, expiring OTP challenge"),
            ("3", "Generate website", "Upload media, generate AI content and store doctor record"),
            ("4", "Publish doctor portal", "Expose public doctor route using a unique slug"),
            ("5", "Manage operations", "Doctor admin handles appointments, content and chatbot"),
        ],
        widths=[0.55, 2.65, 3.1],
        header_fill=TEAL,
    )
    add_caption("Figure 3.2: Doctor website generation workflow")


# Cover page
section.different_first_page_header_footer = True
section.first_page_header.paragraphs[0].text = ""
section.first_page_footer.paragraphs[0].text = ""
add_center("A MINI PROJECT REPORT", 12, True, RGBColor.from_string(TEAL), after=12)
if LOGO.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture = p.add_run().add_picture(str(LOGO), width=Inches(4.6))
    doc_pr = picture._inline.docPr
    doc_pr.set("descr", "Shreyarth University logo")
    doc_pr.set("title", "Shreyarth University")
    p.paragraph_format.space_after = Pt(12)
add_center("CARE-GRID", 30, True, RGBColor.from_string(NAVY), after=2)
add_center("(Doctor Website Maker)", 16, True, RGBColor.from_string(TEAL), after=20)
add_center("Submitted in partial fulfilment of the requirements for the degree of", 10.5, False, DARK, after=4)
add_center("Integrated B.Sc. - M.Sc. (CA & IT)", 12, True, DARK, after=20)
add_center("Submitted by", 10, True, MUTED, after=3)
add_center("PARTH KADIYA", 15, True, RGBColor.from_string(NAVY), after=2)
add_center("Enrollment No.: 224005013", 11, True, DARK, after=20)
add_center("School of Computer Science and Application", 11, True, DARK, after=2)
add_center("Shreyarth University, Ahmedabad", 12, True, RGBColor.from_string(NAVY), after=2)
add_center("Academic Year 2025-26 | May 2026", 10.5, False, MUTED, after=0)

# Certificate
doc.add_page_break()
add_center("CERTIFICATE", 20, True, RGBColor.from_string(NAVY), after=18)
add_para(
    "This is to certify that the project report entitled “Care-Grid (Doctor Website Maker)” has been carried out by Parth Kadiya "
    "(Enrollment No. 224005013) under the prescribed academic framework in partial fulfilment of the requirements for the Integrated "
    "Master of Science in CA & IT at the School of Computer Science and Application, Shreyarth University, Ahmedabad, during the academic year 2025-26."
)
add_para(
    "The work presented in this report is a bonafide record of the design, implementation, testing and production-readiness activities "
    "performed for the Care-Grid platform."
)
doc.add_paragraph("\n\n\n")
add_table(
    ["Internal Guide", "Head of Department"],
    [("Name: ____________________\nSignature: ________________", "Name: ____________________\nSignature: ________________")],
    widths=[3.15, 3.15],
    header_fill=NAVY,
)
add_center("School of Computer Science and Application\nShreyarth University, Ahmedabad", 10, True, MUTED, after=0)

# Declaration
doc.add_page_break()
add_center("DECLARATION", 20, True, RGBColor.from_string(NAVY), after=18)
add_para(
    "I hereby declare that the project report entitled “Care-Grid (Doctor Website Maker)”, submitted in partial fulfilment of the "
    "Integrated Master of Science in CA & IT to Shreyarth University, Ahmedabad, is an original record of work carried out by me. "
    "The report has not been submitted to any other university or institution for the award of any degree or diploma."
)
add_para(
    "All technologies, documentation and external resources used during the project have been appropriately acknowledged. "
    "No confidential keys, passwords or deployment secrets are included in this report."
)
doc.add_paragraph("\n\n\n")
add_table(
    ["Student Name", "Enrollment No.", "Signature"],
    [("Parth Kadiya", "224005013", "____________________")],
    widths=[2.3, 2.0, 2.0],
    header_fill=TEAL,
)

# Acknowledgement
doc.add_page_break()
add_center("ACKNOWLEDGEMENT", 20, True, RGBColor.from_string(NAVY), after=18)
add_para(
    "I express my sincere gratitude to the faculty members of the School of Computer Science and Application, Shreyarth University, "
    "for providing the academic foundation, guidance and encouragement required to complete this project. Their feedback helped shape "
    "the project from an initial doctor website generator into a secure, database-driven healthcare portal."
)
add_para(
    "I am thankful to my internal guide and the Head of Department for their support throughout the analysis, development, testing and "
    "documentation phases. I also acknowledge my friends and family for their patience and motivation during the project."
)
add_para(
    "Finally, I acknowledge the open-source communities and official documentation of React, Node.js, Express, MongoDB, Cloudinary, "
    "Nodemailer, OpenAI, Vercel and Render, which supported the implementation and production-readiness of Care-Grid."
)
add_center("Parth Kadiya | Enrollment No. 224005013", 10, True, MUTED, after=0)

# Abstract
doc.add_page_break()
add_center("ABSTRACT", 20, True, RGBColor.from_string(NAVY), after=18)
add_para(
    "Care-Grid is a full-stack MERN platform that enables a doctor to create a responsive professional website by completing a structured "
    "onboarding form. The platform generates public-facing website content, manages profile and clinic gallery images, exposes appointment "
    "booking to patients, and provides dedicated operational dashboards for doctors and a centralized super administrator."
)
add_para(
    "The system uses React and Vite for the frontend, Node.js and Express for the API layer, MongoDB for persistent storage, Cloudinary for "
    "media management, Gmail SMTP for one-time-password delivery, and OpenAI services for grounded administrative assistance and website content. "
    "Appointments submitted from a doctor's public website are stored in MongoDB and become visible in both the doctor dashboard and the super-admin dashboard."
)
add_para(
    "Security is treated as a core requirement. Registration, doctor login and super-admin login use purpose-bound email OTP verification. "
    "Passwords and OTP values are stored as bcrypt hashes, access tokens are role-bound, password reset invalidates older doctor sessions, API routes "
    "are protected with authorization middleware, and validation and rate limiting are enforced across sensitive workflows."
)
add_callout(
    "Project Outcome",
    "Care-Grid provides a practical, production-oriented foundation for creating and managing doctor websites, appointments and healthcare portal operations from a single platform.",
)

# TOC
doc.add_page_break()
add_center("TABLE OF CONTENTS", 20, True, RGBColor.from_string(NAVY), after=18)
add_table(
    ["Section", "Title", "Page"],
    [
        ("Front Matter", "Certificate", "2"),
        ("Front Matter", "Declaration", "3"),
        ("Front Matter", "Acknowledgement", "4"),
        ("Front Matter", "Abstract", "5"),
        ("Front Matter", "List of Figures and Tables", "7"),
        ("Chapter 1", "Introduction and Project Overview", "8"),
        ("Chapter 2", "Requirements and System Analysis", "11"),
        ("Chapter 3", "System Design, Architecture and Database", "14"),
        ("Chapter 4", "Implementation", "17"),
        ("Chapter 5", "Security, Validation and Error Handling", "20"),
        ("Chapter 6", "Testing and Quality Assurance", "23"),
        ("Chapter 7", "Deployment, Conclusion and Future Scope", "26"),
        ("References", "Bibliography and References", "28"),
        ("Appendix", "Project Structure and User Guide", "29"),
    ],
    widths=[1.35, 4.3, 0.65],
    header_fill=NAVY,
)

# Lists
doc.add_page_break()
add_center("LIST OF FIGURES AND TABLES", 20, True, RGBColor.from_string(NAVY), after=18)
add_table(
    ["Type", "Identifier", "Title"],
    [
        ("Figure", "3.1", "High-level Care-Grid system architecture"),
        ("Figure", "3.2", "Doctor website generation workflow"),
        ("Table", "2.1", "Functional requirements"),
        ("Table", "2.2", "Non-functional requirements"),
        ("Table", "3.1", "Database collections"),
        ("Table", "4.1", "Core application routes"),
        ("Table", "5.1", "Security controls"),
        ("Table", "6.1", "Testing summary"),
    ],
    widths=[1.0, 1.0, 4.3],
    header_fill=NAVY,
)

# Chapter 1
add_chapter("Chapter 1", "Introduction and Project Overview")
doc.add_heading("1.1 Introduction", level=1)
add_para(
    "Doctors and small clinics increasingly require a trustworthy digital presence, but building and maintaining an individual website, appointment workflow "
    "and administrative dashboard generally requires technical expertise and multiple disconnected services. Care-Grid addresses this challenge by offering a "
    "single platform that converts structured doctor and clinic information into a modern public website and an operational management environment."
)
add_para(
    "A doctor enters professional details, clinic information, contact information, working schedule, website theme, profile image and gallery images. "
    "After verifying the submitted email address using an OTP, Care-Grid stores the information and publishes a unique doctor website. Patients can then "
    "view treatments, clinic details, gallery images and directions, and can submit appointment requests."
)
doc.add_heading("1.2 Problem Statement", level=2)
add_bullets([
    "Independent doctors may not have the technical resources to create and maintain a responsive website.",
    "Patient appointment requests are often collected through unstructured calls or messages and are difficult to track.",
    "A doctor requires a private dashboard while a platform owner requires a centralized view across all doctor portals.",
    "Authentication and account recovery must be secure without making the workflow difficult for non-technical users.",
    "Public websites, media, administrative data and AI-powered assistance need to work together without exposing sensitive credentials or ungrounded information.",
])
doc.add_heading("1.3 Project Objectives", level=2)
add_numbers([
    "Enable doctors to generate a professional website through a guided form.",
    "Require verified email ownership before creating a doctor portal.",
    "Store doctor, appointment, OTP and chatbot records in MongoDB.",
    "Provide role-specific doctor and super-admin dashboards with analytics.",
    "Support appointment booking and status management.",
    "Provide a data-grounded chatbot that responds from admin-panel data.",
    "Apply responsive design, validation, secure authentication and production deployment practices.",
])
doc.add_heading("1.4 Scope", level=2)
add_para(
    "The current scope covers doctor website generation, responsive themes, profile and gallery media, public appointment booking, doctor dashboard operations, "
    "super-admin monitoring, grounded chatbot support, OTP-based authentication, password recovery, deployment configuration and health monitoring. "
    "The platform does not provide medical diagnosis, online prescriptions, payment processing or a full hospital information management system."
)
doc.add_heading("1.5 Key Stakeholders", level=2)
add_table(
    ["Stakeholder", "Primary needs"],
    [
        ("Doctor", "Create and update website; view and manage appointments; access grounded dashboard assistance"),
        ("Patient", "View trusted clinic information; request an appointment; open clinic directions"),
        ("Super Admin", "Monitor all doctors and appointments; edit/remove portals; use network-level chatbot"),
        ("System Administrator", "Deploy, secure, monitor and maintain the application and database"),
    ],
    widths=[1.6, 4.7],
    header_fill=TEAL,
)

# Chapter 2
add_chapter("Chapter 2", "Requirements and System Analysis")
doc.add_heading("2.1 Existing System Analysis", level=1)
add_para(
    "Traditional clinic websites are commonly static and require a developer for every update. Appointment requests may be handled through phone calls, "
    "social media messages or generic forms without a unified administrative record. These approaches create fragmented information, inconsistent branding, "
    "limited visibility and weak operational tracking."
)
doc.add_heading("2.2 Proposed System", level=2)
add_para(
    "Care-Grid proposes a multi-portal platform where the same application can generate and operate multiple doctor websites. Each doctor receives a unique "
    "slug-based public route and a protected dashboard. A super administrator receives an aggregated view of doctor portals and patient appointment requests."
)
doc.add_heading("2.3 Functional Requirements", level=2)
add_table(
    ["ID", "Requirement", "Priority"],
    [
        ("FR-01", "Doctor submits verified details and creates a website", "High"),
        ("FR-02", "System generates a unique public doctor URL", "High"),
        ("FR-03", "Patient submits appointment request", "High"),
        ("FR-04", "Doctor views and updates appointment status", "High"),
        ("FR-05", "Super admin views doctors and appointments", "High"),
        ("FR-06", "Doctor updates profile, content, theme and media", "Medium"),
        ("FR-07", "Doctor and super admin use data-grounded chatbot", "Medium"),
        ("FR-08", "OTP verifies registration, login and password recovery", "High"),
        ("FR-09", "Google Maps searches clinic name and address", "Medium"),
    ],
    widths=[0.75, 4.75, 0.8],
    header_fill=NAVY,
)
add_caption("Table 2.1: Functional requirements")
doc.add_heading("2.4 Non-Functional Requirements", level=2)
add_table(
    ["Category", "Requirement"],
    [
        ("Security", "Hash passwords and OTPs; enforce role authorization; apply CORS, Helmet, validation and rate limiting"),
        ("Usability", "Provide responsive interfaces, custom feedback messages and simple guided flows"),
        ("Reliability", "Persist records in MongoDB and preserve existing Cloudinary assets when replacement fails"),
        ("Performance", "Use indexed database fields, bounded payloads and limited chatbot context"),
        ("Maintainability", "Organize code into routes, controllers, services, models, validations and reusable React components"),
        ("Deployability", "Support Vercel frontend, Render backend and MongoDB Atlas with environment-based configuration"),
    ],
    widths=[1.45, 4.85],
    header_fill=TEAL,
)
add_caption("Table 2.2: Non-functional requirements")
doc.add_heading("2.5 Feasibility Study", level=2)
add_bullets([
    "Technical feasibility: The MERN stack and selected managed services support all required workflows.",
    "Operational feasibility: Guided forms and dashboards reduce the need for specialized technical skills.",
    "Economic feasibility: Vercel, Render, MongoDB Atlas and Cloudinary offer entry-level managed hosting options.",
    "Schedule feasibility: Modular development enables features to be implemented and tested incrementally.",
])

# Chapter 3
add_chapter("Chapter 3", "System Design, Architecture and Database")
doc.add_heading("3.1 Architecture", level=1)
add_para(
    "Care-Grid follows a three-layer web architecture. The React client handles presentation and interactions. The Express API applies validation, business logic, "
    "authentication and integrations. MongoDB and external services persist data and provide media, email and AI capabilities."
)
add_architecture_figure()
doc.add_heading("3.2 Doctor Website Generation Workflow", level=2)
add_workflow_figure()
doc.add_heading("3.3 Database Design", level=2)
add_para(
    "MongoDB is used because the platform stores nested doctor profiles, flexible AI-generated content, appointment snapshots, OTP metadata and chat message arrays. "
    "Mongoose schemas enforce required fields, allowed values, indexes and relationships."
)
add_table(
    ["Collection", "Important fields", "Purpose"],
    [
        ("Doctor", "slug, theme, password, authVersion, personalDetails, media, aiContent", "Stores doctor portal and protected account data"),
        ("Appointment", "doctor, doctorSnapshot, patient, treatment, date, timeSlot, status", "Stores patient appointment requests"),
        ("Admin", "email, password", "Stores super-admin credentials"),
        ("OtpChallenge", "accountType, purpose, accountId, otpHash, expiresAt, attempts", "Stores purpose-bound one-time verification challenges"),
        ("ChatConversation", "ownerType, owner, messages", "Persists doctor and super-admin chatbot history"),
    ],
    widths=[1.25, 3.0, 2.05],
    header_fill=NAVY,
)
add_caption("Table 3.1: MongoDB collections")
doc.add_heading("3.4 Key Relationships and Indexes", level=2)
add_bullets([
    "Each appointment references one doctor using the doctor ObjectId.",
    "Chat conversations use a unique ownerType and owner combination.",
    "Doctor slugs and doctor email addresses are unique.",
    "Appointment doctor/date/status fields are indexed for administrative queries.",
    "OTP expiry uses a TTL index so expired challenges are automatically removed.",
])
doc.add_heading("3.5 External Service Design", level=2)
add_table(
    ["Service", "Use in Care-Grid"],
    [
        ("Cloudinary", "Secure profile and clinic gallery media hosting"),
        ("Gmail SMTP / Nodemailer", "OTP and password-change notification email delivery"),
        ("OpenAI API", "Doctor website content and grounded administrative assistance"),
        ("Google Maps Search", "Clinic directions using clinic name, address and country"),
        ("MongoDB Atlas", "Managed production database"),
    ],
    widths=[1.65, 4.65],
    header_fill=TEAL,
)

# Chapter 4
add_chapter("Chapter 4", "Implementation")
doc.add_heading("4.1 Technology Stack", level=1)
add_table(
    ["Layer", "Technology", "Responsibility"],
    [
        ("Frontend", "React 19, Vite, Tailwind CSS, Axios", "Responsive pages, components, forms and API communication"),
        ("Backend", "Node.js, Express 5", "REST API, authentication, validation and integrations"),
        ("Database", "MongoDB, Mongoose", "Persistent structured application data"),
        ("Authentication", "JWT, bcrypt, email OTP", "Role access, password hashing and identity verification"),
        ("Operations", "Vercel, Render, MongoDB Atlas", "Production hosting and managed database"),
    ],
    widths=[1.15, 2.2, 2.95],
    header_fill=NAVY,
)
doc.add_heading("4.2 Frontend Modules", level=2)
add_bullets([
    "Home: doctor onboarding form, theme selection, image previews, discard controls and registration OTP.",
    "Doctor Profile: public responsive doctor website, gallery, treatments, contact details, maps and appointment modal.",
    "Doctor Admin Panel: secure login, OTP, forgot password, analytics, appointments, website editor and chatbot.",
    "Super Admin Login and Dashboard: secure OTP login, doctor directory, appointments, analytics, editing and chatbot.",
    "Reusable components: OTP verification, file upload, gallery slider, chatbot, appointment manager, analytics and UI provider.",
])
doc.add_heading("4.3 Backend Modules", level=2)
add_bullets([
    "Routes define public, doctor-protected and super-admin-protected endpoints.",
    "Controllers implement doctor, appointment, authentication, admin and chatbot workflows.",
    "Services isolate OTP, SMTP, Cloudinary, OpenAI and grounded chatbot behavior.",
    "Mongoose models define persistent data structures and indexes.",
    "Joi validations reject malformed or unsafe request data before controller logic.",
    "Middleware handles authentication, upload validation, error responses, logging and authorization.",
])
doc.add_heading("4.4 Core Application Routes", level=2)
add_table(
    ["Method", "Route", "Purpose"],
    [
        ("POST", "/api/auth/registration-otp", "Send registration verification code"),
        ("POST", "/api/auth/verify-otp", "Verify purpose-bound OTP"),
        ("POST", "/api/doctors/register", "Create verified doctor portal"),
        ("GET", "/api/doctors/:slug", "Load public doctor website"),
        ("POST", "/api/doctors/:slug/appointments", "Create patient appointment"),
        ("POST", "/api/doctors/:slug/login", "Verify doctor credentials and send OTP"),
        ("GET", "/api/doctors/:slug/session", "Validate doctor session and load profile"),
        ("GET", "/api/admin/doctors", "Load all doctor portals for super admin"),
        ("GET", "/api/health", "Report backend and database health"),
    ],
    widths=[0.8, 2.65, 2.85],
    header_fill=TEAL,
)
add_caption("Table 4.1: Core application routes")
doc.add_heading("4.5 OTP Authentication Flow", level=2)
add_numbers([
    "The user submits valid credentials or a registration email.",
    "The backend generates a cryptographically random six-digit OTP.",
    "Only the bcrypt hash of the OTP is saved with purpose, expiry and attempt count.",
    "The OTP is delivered through Gmail SMTP.",
    "The user enters the OTP using the verification interface.",
    "The backend verifies purpose, expiry, attempts and hash, then deletes the challenge.",
    "A purpose-specific short-lived token or role access token is issued.",
])
doc.add_heading("4.6 Example Configuration", level=2)
add_code(
    "Frontend environment:\n"
    "VITE_API_URL=https://care-grid-api.onrender.com\n\n"
    "Backend deployment:\n"
    "Root Directory: server\n"
    "Build Command: npm ci --omit=dev\n"
    "Start Command: npm start\n"
    "Health Check Path: /api/health"
)

# Chapter 5
add_chapter("Chapter 5", "Security, Validation and Error Handling")
doc.add_heading("5.1 Security Approach", level=1)
add_para(
    "Care-Grid applies defense in depth. No individual control is treated as sufficient. Input validation, authentication, authorization, secure hashing, "
    "rate limiting, origin restrictions, payload limits and production secret management work together to reduce risk."
)
add_table(
    ["Control", "Implementation", "Risk reduced"],
    [
        ("Password storage", "bcrypt hashing", "Plaintext credential exposure"),
        ("OTP storage", "Purpose-bound bcrypt hashes, expiry and attempt limit", "OTP theft, replay and cross-flow misuse"),
        ("Authorization", "JWT roles and protected middleware", "Unauthorized dashboard access"),
        ("Session invalidation", "Doctor authVersion increment after reset", "Continued use of old sessions"),
        ("Validation", "Joi schemas and Mongoose rules", "Malformed and inconsistent records"),
        ("Rate limiting", "Global and sensitive-route limits", "Brute force and request abuse"),
        ("HTTP security", "Helmet, CORS allowlist and hidden framework signature", "Common browser and origin attacks"),
        ("Upload security", "Size, count, MIME and magic-byte validation", "Spoofed or oversized file uploads"),
        ("Secret management", "Environment variables and Git ignore rules", "Credential leakage"),
    ],
    widths=[1.45, 2.85, 2.0],
    header_fill=NAVY,
)
add_caption("Table 5.1: Security controls")
doc.add_heading("5.2 Validation Rules", level=2)
add_bullets([
    "Doctor mobile number must contain exactly ten digits.",
    "Doctor registration email must be verified and unique.",
    "Passwords require uppercase, lowercase, number and special character.",
    "Appointment status is restricted to pending, confirmed, completed or cancelled.",
    "Uploaded media is restricted to JPEG, PNG or WebP and maximum file/count limits.",
    "Chat messages and history have bounded length.",
])
doc.add_heading("5.3 Error Handling", level=2)
add_para(
    "The API uses a centralized error middleware to produce consistent JSON responses. Expected validation and authentication errors return client-appropriate "
    "status codes. Unexpected errors are logged without returning stack traces or secrets. The frontend displays custom feedback rather than browser alert boxes."
)
add_table(
    ["Condition", "HTTP status", "Response behavior"],
    [
        ("Invalid form or JSON", "400", "Return concise validation message"),
        ("Missing or invalid session", "401", "Reject request and require login"),
        ("Wrong role or CORS origin", "403", "Reject unauthorized access"),
        ("Missing record or route", "404", "Return not-found response"),
        ("Duplicate unique record", "409", "Return conflict response"),
        ("Service temporarily unavailable", "503", "Return safe retry message"),
    ],
    widths=[2.25, 1.15, 2.9],
    header_fill=TEAL,
)
doc.add_heading("5.4 Privacy Considerations", level=2)
add_para(
    "The administrative chatbot receives a minimized data context. Patient phone numbers, patient email addresses and free-form patient messages are not included "
    "in the OpenAI context. The chatbot is instructed to answer only from administrative data and not provide medical diagnosis or treatment advice."
)

# Chapter 6
add_chapter("Chapter 6", "Testing and Quality Assurance")
doc.add_heading("6.1 Testing Strategy", level=1)
add_para(
    "The project was tested through static checks, production builds, dependency audits, schema validation tests, live API regression checks, database audits "
    "and production-mode startup verification. Testing focused on both expected workflows and invalid or malicious inputs."
)
add_table(
    ["Test area", "Representative check", "Result"],
    [
        ("Frontend quality", "ESLint across React source", "Passed"),
        ("Frontend production", "Vite production build", "Passed"),
        ("Backend syntax", "Syntax check for all backend JavaScript files", "Passed"),
        ("Dependency security", "npm audit --omit=dev for client and server", "0 vulnerabilities"),
        ("Health check", "GET /api/health with connected database", "200 healthy"),
        ("Authorization", "Unauthenticated admin request", "401 rejected"),
        ("CORS", "Request from unauthorized origin", "403 rejected"),
        ("Validation", "Eleven-digit doctor mobile", "400 rejected"),
        ("Input parsing", "Malformed JSON request", "400 rejected"),
        ("Upload security", "Spoofed JPEG file", "400 rejected"),
        ("Database", "Unique doctor email index and record scan", "Passed"),
        ("Production startup", "Render-style NODE_ENV=production smoke test", "Passed"),
    ],
    widths=[1.5, 3.5, 1.3],
    header_fill=NAVY,
)
add_caption("Table 6.1: Testing summary")
doc.add_heading("6.2 Sample Test Cases", level=2)
add_table(
    ["ID", "Scenario", "Expected result"],
    [
        ("TC-01", "Doctor registration without OTP token", "Registration rejected"),
        ("TC-02", "Reuse a successfully verified OTP", "Second verification rejected"),
        ("TC-03", "Doctor login with correct password", "OTP sent before dashboard access"),
        ("TC-04", "Password reset followed by old JWT use", "Old session rejected"),
        ("TC-05", "Patient books appointment", "Record visible to doctor and super admin"),
        ("TC-06", "Doctor opens maps direction", "Clinic name and address searched together"),
        ("TC-07", "Replace image upload fails", "Existing image remains available"),
    ],
    widths=[0.75, 3.5, 2.05],
    header_fill=TEAL,
)
doc.add_heading("6.3 Database Quality Findings", level=2)
add_para(
    "The production-readiness audit confirmed that doctor email addresses have a unique database index, OTP challenge storage was empty after testing, "
    "and appointment and doctor collections were accessible. A legacy doctor record containing a twelve-digit phone number was identified; the system now "
    "rejects such values, and the legacy record must be corrected using the actual ten-digit number rather than unsafe automatic truncation."
)

# Chapter 7
add_chapter("Chapter 7", "Deployment, Conclusion and Future Scope")
doc.add_heading("7.1 Production Deployment Design", level=1)
add_para(
    "Care-Grid is prepared for a split deployment model. The React/Vite frontend is deployed on Vercel, the Node/Express backend is deployed on Render, "
    "and MongoDB Atlas provides the production database. Cloudinary, Gmail SMTP and OpenAI remain managed external integrations."
)
add_table(
    ["Platform", "Configuration"],
    [
        ("GitHub", "Source repository with secrets, node_modules, logs and builds excluded"),
        ("Vercel", "Root directory client; VITE_API_URL points to Render service"),
        ("Render", "Root directory server; npm ci --omit=dev; npm start; /api/health"),
        ("MongoDB Atlas", "SRV connection string, dedicated user and controlled network access"),
    ],
    widths=[1.35, 4.95],
    header_fill=NAVY,
)
doc.add_heading("7.2 Required Production Environment", level=2)
add_code(
    "NODE_ENV=production\n"
    "MONGO_URI=<MongoDB Atlas SRV URI>\n"
    "CLIENT_URLS=<Vercel production URL>\n"
    "JWT_SECRET=<long random secret>\n"
    "OPENAI_API_KEY=<secret>\n"
    "CLOUDINARY_CLOUD_NAME / CLOUDINARY_API_KEY / CLOUDINARY_API_SECRET\n"
    "EMAIL_USER / EMAIL_APP_PASSWORD\n"
    "ADMIN_EMAIL / ADMIN_PASSWORD (initial bootstrap only)"
)
doc.add_heading("7.3 Conclusion", level=2)
add_para(
    "Care-Grid successfully demonstrates how a modern MERN application can simplify doctor website creation while supporting real operational workflows. "
    "The platform connects public doctor websites, patient appointment booking, doctor administration and centralized supervision through a single database-driven system."
)
add_para(
    "The completed implementation goes beyond a basic website generator by integrating secure OTP authentication, password recovery, responsive themes, media management, "
    "grounded AI assistance, analytics dashboards, persistent chat history, validation, production deployment configuration and security hardening."
)
doc.add_heading("7.4 Future Scope", level=2)
add_bullets([
    "Add doctor subscription plans, billing and custom domains.",
    "Add verified patient accounts and appointment reminders.",
    "Add calendar integration and precise time-slot availability.",
    "Add multilingual doctor websites and accessibility enhancements.",
    "Add audit logs, admin activity history and advanced reporting.",
    "Add secure teleconsultation while maintaining healthcare privacy requirements.",
    "Add automated database backups, uptime monitoring and incident alerts.",
])

# Bibliography
doc.add_page_break()
doc.add_heading("BIBLIOGRAPHY AND REFERENCES", level=1)
references = [
    "React Documentation. https://react.dev/",
    "Vite Documentation. https://vite.dev/",
    "Node.js Documentation. https://nodejs.org/docs/latest/api/",
    "Express Documentation. https://expressjs.com/",
    "MongoDB and Mongoose Documentation. https://www.mongodb.com/docs/ and https://mongoosejs.com/docs/",
    "OWASP Cheat Sheet Series: Authentication, Password Storage and Forgot Password. https://cheatsheetseries.owasp.org/",
    "Nodemailer SMTP Documentation. https://nodemailer.com/smtp",
    "Cloudinary Documentation. https://cloudinary.com/documentation",
    "OpenAI API Documentation. https://platform.openai.com/docs/",
    "Vercel Vite Deployment Documentation. https://vercel.com/docs/frameworks/vite",
    "Render Node/Express Deployment Documentation. https://render.com/docs/deploy-node-express-app",
    "MongoDB Atlas Documentation. https://www.mongodb.com/docs/atlas/",
]
add_numbers(references)

# Appendix
doc.add_page_break()
doc.add_heading("APPENDIX A: PROJECT DIRECTORY STRUCTURE", level=1)
add_code(
    "doctor-portal-maker/\n"
    "  client/\n"
    "    src/components/       Reusable UI components\n"
    "    src/pages/            Public, doctor-admin and super-admin pages\n"
    "    src/utils/            Client authentication utilities\n"
    "    vercel.json           SPA rewrite configuration\n"
    "  server/\n"
    "    controllers/          Business workflow handlers\n"
    "    middlewares/          Authentication, validation, uploads and errors\n"
    "    models/               MongoDB/Mongoose schemas\n"
    "    routes/               REST API routes\n"
    "    services/             Email, OTP, Cloudinary and AI integrations\n"
    "    validations/          Joi request schemas\n"
    "  render.yaml             Render deployment blueprint\n"
    "  PRODUCTION_DEPLOYMENT.md"
)
doc.add_heading("APPENDIX B: USER GUIDE", level=1)
doc.add_heading("For Doctors", level=2)
add_numbers([
    "Open the Care-Grid website creation form and enter valid professional and clinic information.",
    "Upload a profile image and optional clinic gallery images.",
    "Verify the submitted email using the received OTP.",
    "Open the generated public website and doctor admin URL.",
    "Sign in using registered credentials and complete OTP verification.",
    "Use the dashboard to manage appointments, website content, media, theme and chatbot.",
])
doc.add_heading("For Patients", level=2)
add_numbers([
    "Open the doctor's public Care-Grid website.",
    "Review doctor, clinic, treatments, gallery and directions information.",
    "Select Book Appointment and submit the required information.",
    "Wait for the clinic to confirm or update the appointment request.",
])
doc.add_heading("For Super Admin", level=2)
add_numbers([
    "Open /superadmin and sign in using the authorized account.",
    "Complete email OTP verification.",
    "Review all doctors, appointments and network analytics.",
    "Edit or remove doctor portals when authorized.",
    "Use the grounded chatbot for administrative questions.",
])

# Set core properties and save
doc.core_properties.title = "Care-Grid (Doctor Website Maker) - Project Report"
doc.core_properties.subject = "MERN Stack Doctor Website Maker"
doc.core_properties.author = "Parth Kadiya"
doc.core_properties.keywords = "Care-Grid, MERN, Doctor Website, MongoDB, React, Express, OTP"
doc.core_properties.comments = "Prepared for Shreyarth University academic submission."

for sec in doc.sections:
    configure_header_footer(sec)

for table in doc.tables:
    if table.rows:
        repeat_table_header(table.rows[0])

doc.save(OUTPUT)
print(OUTPUT)
